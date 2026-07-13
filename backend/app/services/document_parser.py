"""
Local document parsing service.
Supports PDF (PyMuPDF + Tesseract OCR fallback), DOCX, TXT.
Returns structured ParsedDocument with page text, metadata, and parsing warnings.
No external calls — fully offline.
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

from app.core.logging import get_logger

log = get_logger(__name__)


@dataclass
class PageContent:
    page_number: int
    text: str
    has_images: bool = False
    has_tables: bool = False


@dataclass
class ParsedDocument:
    raw_text: str
    pages: List[PageContent] = field(default_factory=list)
    page_count: int = 0
    metadata: dict = field(default_factory=dict)
    parsing_warnings: List[str] = field(default_factory=list)
    used_ocr: bool = False
    file_extension: str = ""


# ── Parsing risk detectors ────────────────────────────────────────────────────

def _detect_multi_column(text: str) -> bool:
    """Heuristic: many very short lines suggests multi-column layout."""
    lines = [l for l in text.splitlines() if l.strip()]
    if not lines:
        return False
    short = sum(1 for l in lines if len(l.strip()) < 40)
    return short / len(lines) > 0.6


def _detect_missing_standard_sections(text: str) -> List[str]:
    """Return list of common ATS sections that appear to be missing."""
    lower = text.lower()
    missing = []
    for section, patterns in {
        "Contact Information": ["email", "phone", "@"],
        "Work Experience": ["experience", "employment", "work history", "professional"],
        "Education": ["education", "degree", "university", "college", "bachelor", "master"],
        "Skills": ["skills", "technical", "proficiency"],
    }.items():
        if not any(p in lower for p in patterns):
            missing.append(section)
    return missing


def _detect_excessive_graphics(text: str, has_images: bool) -> bool:
    if has_images and len(text.strip()) < 200:
        return True
    return False


def _assess_parsing_risks(parsed: ParsedDocument) -> List[str]:
    warnings: List[str] = []
    if _detect_multi_column(parsed.raw_text):
        warnings.append("Multi-column layout detected — ATS parsers may misread column order.")
    if parsed.used_ocr:
        warnings.append("Scanned or image-based PDF detected — OCR used; accuracy may vary.")
    if any(p.has_images for p in parsed.pages):
        warnings.append("Images found in document — text inside images is not extracted unless OCR was applied.")
    if any(p.has_tables for p in parsed.pages):
        warnings.append("Tables detected — ATS systems may not parse table content correctly.")
    missing_sections = _detect_missing_standard_sections(parsed.raw_text)
    for s in missing_sections:
        warnings.append(f"Standard section possibly missing: {s}")
    if parsed.raw_text and "http" not in parsed.raw_text.lower() and "@" not in parsed.raw_text:
        warnings.append("No email address detected in document — contact information may be incomplete.")
    if len(parsed.raw_text.strip()) < 100:
        warnings.append("Very little text extracted — document may be image-only or empty.")
    return warnings


# ── PDF parser ────────────────────────────────────────────────────────────────

def parse_pdf(data: bytes) -> ParsedDocument:
    """Parse PDF using PyMuPDF; fallback to Tesseract for scanned pages."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise RuntimeError("PyMuPDF (fitz) is not installed.")

    doc = fitz.open(stream=data, filetype="pdf")
    pages: List[PageContent] = []
    all_text_parts: List[str] = []
    used_ocr = False

    for page_num, page in enumerate(doc, start=1):
        text = page.get_text("text")
        has_images = len(page.get_images(full=True)) > 0
        has_tables = False

        # Detect tables via block analysis
        blocks = page.get_text("blocks")
        # Simple heuristic: if >3 blocks on same Y-level → likely table
        y_positions = [round(b[1]) for b in blocks if b[6] == 0]
        y_counts = {}
        for y in y_positions:
            y_counts[y] = y_counts.get(y, 0) + 1
        if any(v > 2 for v in y_counts.values()):
            has_tables = True

        # OCR fallback for pages with minimal text but images
        if len(text.strip()) < 50 and has_images:
            try:
                import pytesseract
                from PIL import Image

                pix = page.get_pixmap(dpi=200)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                ocr_text = pytesseract.image_to_string(img, lang="eng")
                if len(ocr_text.strip()) > len(text.strip()):
                    text = ocr_text
                    used_ocr = True
            except Exception as e:
                log.warning("OCR fallback failed for page %d: %s", page_num, e)

        pages.append(PageContent(
            page_number=page_num,
            text=text,
            has_images=has_images,
            has_tables=has_tables,
        ))
        all_text_parts.append(text)

    raw_text = "\n".join(all_text_parts)

    # Metadata
    metadata = {}
    try:
        meta = doc.metadata
        if meta:
            metadata = {k: v for k, v in meta.items() if v and k in ("title", "author", "creator", "producer")}
    except Exception:
        pass
    doc.close()

    parsed = ParsedDocument(
        raw_text=raw_text,
        pages=pages,
        page_count=len(pages),
        metadata=metadata,
        used_ocr=used_ocr,
        file_extension=".pdf",
    )
    parsed.parsing_warnings = _assess_parsing_risks(parsed)
    return parsed


# ── DOCX parser ───────────────────────────────────────────────────────────────

def parse_docx(data: bytes) -> ParsedDocument:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise RuntimeError("python-docx is not installed.")

    doc = DocxDocument(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs]
    # Extract table text too
    table_texts = []
    has_tables = len(doc.tables) > 0
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                table_texts.append(row_text)

    all_text = "\n".join(paragraphs + table_texts)
    pages = [PageContent(page_number=1, text=all_text, has_tables=has_tables)]

    parsed = ParsedDocument(
        raw_text=all_text,
        pages=pages,
        page_count=1,
        file_extension=".docx",
    )
    if has_tables:
        parsed.parsing_warnings.append("Tables detected in DOCX — ATS systems may not parse table content correctly.")
    parsed.parsing_warnings += _assess_parsing_risks(parsed)
    return parsed


# ── TXT parser ────────────────────────────────────────────────────────────────

def parse_txt(data: bytes) -> ParsedDocument:
    import chardet

    detected = chardet.detect(data[:4096])
    encoding = detected.get("encoding") or "utf-8"
    try:
        text = data.decode(encoding, errors="replace")
    except Exception:
        text = data.decode("utf-8", errors="replace")

    pages = [PageContent(page_number=1, text=text)]
    parsed = ParsedDocument(
        raw_text=text,
        pages=pages,
        page_count=1,
        file_extension=".txt",
    )
    parsed.parsing_warnings = _assess_parsing_risks(parsed)
    return parsed


# ── Dispatcher ────────────────────────────────────────────────────────────────

def parse_document(data: bytes, file_extension: str) -> ParsedDocument:
    """Route to the correct parser based on file extension."""
    ext = file_extension.lower()
    if ext == ".pdf":
        return parse_pdf(data)
    elif ext == ".docx":
        return parse_docx(data)
    elif ext == ".txt":
        return parse_txt(data)
    else:
        raise ValueError(f"Unsupported file extension for parsing: {ext}")
